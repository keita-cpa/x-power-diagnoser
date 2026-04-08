"""
株式会社FDS 計算書類 整合性分析スクリプト
対象期間: 令和7年1月1日 ~ 令和7年12月31日（第25期）
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import openpyxl
from docx import Document
from datetime import datetime

# ─── 数値パース補助 ──────────────────────────────
def parse_num(val):
    """カンマ・全角ハイフン等を除去して数値変換"""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return int(val)
    s = str(val).replace(',', '').replace('―', '0').replace('－', '0').strip()
    if s in ('', '―', '－', '-'):
        return 0
    try:
        return int(float(s))
    except ValueError:
        return None

def fmt(val):
    """千円単位の数値を3桁区切りで表示"""
    if val is None:
        return 'N/A'
    return f"{val:,}"

# ─── Excel 読み込み ──────────────────────────────
xl_path = '【FDS】計算書類_20251231.xlsx'
wb = openpyxl.load_workbook(xl_path, data_only=True)

# ── BS（千円）──
ws_bs = wb['BS（千円）']
bs_rows = list(ws_bs.iter_rows(values_only=True))

def find_bs_value(keyword, col_idx=6):
    """BSシートのキーワード行を検索して値を返す（列インデックス0始まり）"""
    for row in bs_rows:
        for cell in row:
            if isinstance(cell, str) and keyword in cell:
                return parse_num(row[col_idx])
    return None

bs_tangible_fixed   = parse_num(bs_rows[18][6])   # 有形固定資産 row19 (0-based: 18)
bs_intangible_fixed = parse_num(bs_rows[25][6])   # 無形固定資産 row26
bs_net_assets       = parse_num(bs_rows[33][14])  # 純資産合計  row34
bs_net_assets_check = parse_num(bs_rows[33][6])   # 純資産合計（左列）
bs_total_assets     = parse_num(bs_rows[34][6])   # 資産合計    row35
bs_kaerkaeken       = parse_num(bs_rows[12][6])   # 立替金      row13

# ── PL（千円）──
ws_pl = wb['PL（千円）']
pl_rows = list(ws_pl.iter_rows(values_only=True))
pl_net_income = parse_num(pl_rows[29][7])  # 当期純利益 row30

# ── SS（千円）──
ws_ss = wb['SS（千円）']
ss_rows = list(ws_ss.iter_rows(values_only=True))

# row14: 繰越利益剰余金 当期変動額（当期純利益）
ss_net_income_change = parse_num(ss_rows[13][8])  # row14, col I (0-based: 8)
# row25: 純資産の部合計 当期末残高
ss_net_assets_end    = parse_num(ss_rows[24][8])  # row25, col I

# ─── 附属明細書 読み込み ──────────────────────────
doc_annex = Document('20260403_FDS_附属明細書_20251231.docx')
annex_table = doc_annex.tables[0]  # 有形・無形固定資産明細テーブル

annex_tangible_total   = None
annex_intangible_total = None

for row in annex_table.rows:
    cells = [c.text.strip() for c in row.cells]
    # '計' が区分・資産種類両方の行 → 種類列で判断
    if cells[1] == '計':
        # 有形固定資産の計
        val = parse_num(cells[6])  # 期末帳簿価額列
        if annex_tangible_total is None:
            annex_tangible_total = val
        elif annex_intangible_total is None:
            annex_intangible_total = val

# ─── 個別注記表 読み込み ──────────────────────────
doc_notes = Document('20260403_FDS_個別注記表20251231.docx')

# Table 2: 関連当事者との取引（兄弟会社等）
#   row1: 親会社 SYNASIA HOLDINGS  立替金  期末残高 = 5,061,744
notes_related_kaerkaeken = None
notes_related_company = None

notes_table2 = doc_notes.tables[2]
for row in notes_table2.rows[1:]:  # ヘッダをスキップ
    cells = [c.text.strip() for c in row.cells]
    # '立替金' を含む行を探す
    if '立替金' in cells[6]:  # 科目列
        notes_related_kaerkaeken = parse_num(cells[7])  # 期末残高列
        notes_related_company    = cells[1].replace('\n', ' ')
        break

# 4. 関係会社に対する金銭債権の段落テキストからも取得
notes_related_receivable = None
for para in doc_notes.paragraphs:
    if '短期金銭債権' in para.text:
        import re
        m = re.search(r'([\d,]+)\s*千円', para.text)
        if m:
            notes_related_receivable = parse_num(m.group(1))
        break

# ─── 整合性チェック ───────────────────────────────
checks = []

def add_check(title, val_a, label_a, val_b, label_b, unit='千円', note=''):
    if val_a is None or val_b is None:
        status = '⚠️ データ取得不可'
        diff = None
    elif val_a == val_b:
        status = '✅ 一致'
        diff = 0
    else:
        diff = val_a - val_b
        status = f'❌ 不一致（差額 {fmt(diff)} {unit}）'
    checks.append({
        'title': title,
        'val_a': val_a, 'label_a': label_a,
        'val_b': val_b, 'label_b': label_b,
        'diff': diff,
        'status': status,
        'unit': unit,
        'note': note,
    })

# Check 1: BS vs SS — 純資産の部合計
add_check(
    '【Check 1】BS vs SS — 純資産の部合計（期末残高）',
    bs_net_assets, 'BS（千円）純資産合計',
    ss_net_assets_end, 'SS（千円）純資産の部合計 当期末残高',
)

# Check 2: PL vs SS — 当期純利益
add_check(
    '【Check 2】PL vs SS — 当期純利益 / 当期変動額（当期純利益）',
    pl_net_income, 'PL（千円）当期純利益',
    ss_net_income_change, 'SS（千円）繰越利益剰余金 当期変動額（当期純利益）',
)

# Check 3: BS vs 附属明細書 — 有形固定資産
add_check(
    '【Check 3】BS vs 附属明細書 — 有形固定資産 期末帳簿価額',
    bs_tangible_fixed, 'BS（千円）有形固定資産',
    annex_tangible_total, '附属明細書 有形固定資産 計（期末帳簿価額）',
)

# Check 4: BS vs 附属明細書 — 無形固定資産
add_check(
    '【Check 4】BS vs 附属明細書 — 無形固定資産 期末帳簿価額',
    bs_intangible_fixed, 'BS（千円）無形固定資産',
    annex_intangible_total, '附属明細書 無形固定資産 計（期末帳簿価額）',
)

# Check 5: BS vs 個別注記表 — 立替金（関連当事者）
add_check(
    '【Check 5】BS vs 個別注記表 — 立替金（関連当事者 SYNASIA HOLDINGS）',
    bs_kaerkaeken, 'BS（千円）立替金（流動資産合計）',
    notes_related_kaerkaeken,
    f'個別注記表 関連当事者取引 立替金 期末残高（{notes_related_company}）',
    note=(
        '※ BS の立替金は全取引先合計。個別注記は関連当事者分のみのため、\n'
        '   差額がある場合は非関連当事者分の立替金が存在する可能性があります。'
    ),
)

# Check 6: BS vs 個別注記表 — 関係会社短期金銭債権
add_check(
    '【Check 6】BS vs 個別注記表 — 関係会社に対する短期金銭債権',
    bs_kaerkaeken, 'BS（千円）立替金（流動資産）',
    notes_related_receivable, '個別注記表 関係会社に対する短期金銭債権',
    note=(
        '※ 個別注記表「貸借対照表に関する注記 4」の関係会社短期金銭債権との照合。'
    ),
)

# ─── Markdown レポート出力 ────────────────────────
report_lines = [
    '# 株式会社FDS 計算書類 整合性分析レポート',
    '',
    f'**分析実施日**: {datetime.today().strftime("%Y年%m月%d日")}  ',
    '**対象期間**: 令和7年1月1日 ～ 令和7年12月31日（第25期）  ',
    '**分析対象ファイル**:',
    '- `【FDS】計算書類_20251231.xlsx`（BS / PL / SS）',
    '- `20260403_FDS_附属明細書_20251231.docx`',
    '- `20260403_FDS_個別注記表20251231.docx`',
    '',
    '---',
    '',
    '## チェック結果サマリー',
    '',
    '| No. | チェック項目 | 結果 |',
    '|-----|-------------|------|',
]

for i, c in enumerate(checks, 1):
    report_lines.append(f'| {i} | {c["title"].replace("【Check " + str(i) + "】", "").strip()} | {c["status"]} |')

report_lines += ['', '---', '', '## 詳細チェック結果', '']

for c in checks:
    report_lines += [
        f'### {c["title"]}',
        '',
        f'| 項目 | 金額（{c["unit"]}） |',
        '|------|------:|',
        f'| {c["label_a"]} | {fmt(c["val_a"])} |',
        f'| {c["label_b"]} | {fmt(c["val_b"])} |',
    ]
    if c['diff'] is not None and c['diff'] != 0:
        report_lines.append(f'| **差額** | **{fmt(c["diff"])}** |')
    report_lines += ['', f'**判定**: {c["status"]}', '']
    if c['note']:
        report_lines += [f'> {c["note"]}', '']
    report_lines.append('')

# ── 主要数値一覧 ──
report_lines += [
    '---',
    '',
    '## 主要数値一覧（BS・PL・SS）',
    '',
    '### 貸借対照表（BS）（単位：千円）',
    '',
    '| 科目 | 金額 |',
    '|------|-----:|',
    f'| 資産合計 | {fmt(bs_total_assets)} |',
    f'| 有形固定資産 | {fmt(bs_tangible_fixed)} |',
    f'| 無形固定資産 | {fmt(bs_intangible_fixed)} |',
    f'| 立替金（流動資産） | {fmt(bs_kaerkaeken)} |',
    f'| 純資産合計 | {fmt(bs_net_assets)} |',
    '',
    '### 損益計算書（PL）（単位：千円）',
    '',
    '| 科目 | 金額 |',
    '|------|-----:|',
    f'| 当期純利益 | {fmt(pl_net_income)} |',
    '',
    '### 株主資本等変動計算書（SS）（単位：千円）',
    '',
    '| 科目 | 金額 |',
    '|------|-----:|',
    f'| 純資産の部合計 当期末残高 | {fmt(ss_net_assets_end)} |',
    f'| 繰越利益剰余金 当期変動額（当期純利益） | {fmt(ss_net_income_change)} |',
    '',
    '---',
    '',
    '*このレポートは analyze_financials.py により自動生成されました。*',
]

report_text = '\n'.join(report_lines)

with open('整合性分析レポート.md', 'w', encoding='utf-8') as f:
    f.write(report_text)

# ─── コンソールサマリー ───────────────────────────
print('=' * 60)
print('  株式会社FDS 計算書類 整合性分析 — コンソールサマリー')
print('=' * 60)
for c in checks:
    print(f'\n{c["title"]}')
    print(f'  {c["label_a"]:<40}: {fmt(c["val_a"])} {c["unit"]}')
    print(f'  {c["label_b"]:<40}: {fmt(c["val_b"])} {c["unit"]}')
    print(f'  → {c["status"]}')
    if c['note']:
        print(f'  {c["note"]}')

print()
print('=' * 60)
print('  レポートを「整合性分析レポート.md」に出力しました。')
print('=' * 60)
